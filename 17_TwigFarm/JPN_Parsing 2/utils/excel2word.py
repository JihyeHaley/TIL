import pandas as pd
from docx import Document
from tqdm import trange
import logging


def excel2word(input_file, output_dir='', log_level=logging.INFO):
    # output_dir의 기입 예시 : /Users/twigfarm/sample/
    logging.basicConfig(format='%(asctime)s :: %(message)s', level=log_level)
    file_name = input_file.split('/')[-1][:-5]
    output_path = f'{output_dir}{file_name}.docx'
    data = pd.read_excel(input_file)

    document = Document()
    # 문서의 기본 정보와 작업자의 정보 기입
    document.add_heading(f'{file_name} 작업 결과물', 0)
    document.add_paragraph('작업자: ', style='List Bullet')
    document.add_paragraph('작업완료일: ', style='List Bullet')
    document.add_paragraph(f'작업한 총 TM 수:      / {len(data)}', style='List Bullet')
    document.add_paragraph(' ')
    # 테이블 생성
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    table.cell(0, 0).text = '#'
    table.cell(0, 1).text = '-' * 110
    logging.info(f'"{file_name}" 파일 처리를 시작합니다.')
    for i in trange(len(data)):
        index = i + 1
        ko, en = data.iloc[i]

        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = str(ko)

        row_cells = table.add_row().cells
        row_cells[1].text = str(en)
    document.save(output_path)
    logging.info(f'{file_name} 파일의 처리가 끝났습니다./n/t/t/t{output_path}를 확인하세요')
